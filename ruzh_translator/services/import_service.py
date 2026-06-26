"""Document import service: parse various file formats into plain text."""

import logging
from pathlib import Path
from typing import Optional

from ruzh_translator.utils.text_utils import (
    detect_language,
    read_text_file,
    split_paragraphs,
    split_sentences,
)

logger = logging.getLogger(__name__)


def import_txt(file_path: str) -> dict:
    """Import a plain text file.

    Args:
        file_path: Path to .txt file.

    Returns:
        Dict with 'raw_text', 'paragraphs', 'language', 'encoding', 'paragraph_count'.
    """
    raw_text = read_text_file(file_path)
    paragraphs = split_paragraphs(raw_text)
    language = detect_language(raw_text)
    return {
        "raw_text": raw_text,
        "paragraphs": paragraphs,
        "language": language,
        "paragraph_count": len(paragraphs),
        "encoding": "utf-8",
    }


def import_docx(file_path: str) -> dict:
    """Import a Word (.docx) file.

    Args:
        file_path: Path to .docx file.

    Returns:
        Dict with 'raw_text', 'paragraphs', 'language', 'paragraph_count'.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    raw_text = "\n\n".join(paragraphs)
    language = detect_language(raw_text)
    return {
        "raw_text": raw_text,
        "paragraphs": paragraphs,
        "language": language,
        "paragraph_count": len(paragraphs),
    }


def import_pdf(file_path: str) -> dict:
    """Import a PDF file.

    Args:
        file_path: Path to .pdf file.

    Returns:
        Dict with 'raw_text', 'paragraphs', 'language', 'paragraph_count'.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    all_text = []
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            all_text.append(text.strip())
    doc.close()

    raw_text = "\n\n".join(all_text)
    paragraphs = split_paragraphs(raw_text)
    language = detect_language(raw_text)
    return {
        "raw_text": raw_text,
        "paragraphs": paragraphs,
        "language": language,
        "paragraph_count": len(paragraphs),
    }


def import_file(file_path: str) -> dict:
    """Import a document file, dispatching to the correct parser.

    Args:
        file_path: Path to the file.

    Returns:
        Dict with import result fields.

    Raises:
        ValueError: If the file format is not supported.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return import_txt(file_path)
    elif suffix == ".docx":
        return import_docx(file_path)
    elif suffix == ".pdf":
        return import_pdf(file_path)
    elif suffix == ".rtf":
        # RTF support via striprtf if available
        try:
            from striprtf.striprtf import rtf_to_text
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                rtf_text = f.read()
            raw_text = rtf_to_text(rtf_text)
            paragraphs = split_paragraphs(raw_text)
            language = detect_language(raw_text)
            return {
                "raw_text": raw_text,
                "paragraphs": paragraphs,
                "language": language,
                "paragraph_count": len(paragraphs),
            }
        except ImportError:
            raise ValueError(
                "RTF support requires striprtf package. Install with: pip install striprtf"
            )
    else:
        raise ValueError(f"Unsupported file format: {suffix}")


def import_and_segment(
    file_path: str,
    session,
    project_id: str,
    language: str = "",
) -> dict:
    """Import a file and automatically split it into sentence segments.

    This is the primary import path for the 'direct translation' workflow.
    It creates both a Document record and Segment records in one step,
    so the translation editor can immediately display the sentences.

    Args:
        file_path: Path to the document file.
        session: SQLAlchemy session.
        project_id: Project ID to associate with.
        language: Language code (auto-detected if empty).

    Returns:
        Dict with 'document', 'segments', 'paragraph_count', 'sentence_count'.
    """
    from ruzh_translator.models.project import Document
    from ruzh_translator.models.segment import Segment

    # Step 1: Parse the file
    data = import_file(file_path)
    language = language or data["language"]
    paragraphs = data["paragraphs"]
    raw_text = data["raw_text"]

    # Step 2: Save the document
    path = Path(file_path)
    doc = Document(
        project_id=project_id,
        filename=path.name,
        file_format=path.suffix.lower(),
        raw_content=raw_text,
        language=language,
        paragraph_count=len(paragraphs),
    )
    session.add(doc)
    session.flush()

    # Step 3: Split into sentence segments
    segments = []
    seg_index = 0
    for para_idx, para_text in enumerate(paragraphs):
        sentences = split_sentences(para_text, language)
        for sent_text in sentences:
            if not sent_text.strip():
                continue
            seg = Segment(
                project_id=project_id,
                document_id=doc.id,
                paragraph_index=para_idx,
                segment_index=seg_index,
                source_text=sent_text.strip(),
                target_text="",
                status="untranslated",
            )
            session.add(seg)
            segments.append(seg)
            seg_index += 1

    session.commit()
    logger.info(
        f"Imported {path.name}: {len(paragraphs)} paragraphs, "
        f"{len(segments)} sentences"
    )

    return {
        "document": doc,
        "segments": segments,
        "paragraph_count": len(paragraphs),
        "sentence_count": len(segments),
        "language": language,
    }
