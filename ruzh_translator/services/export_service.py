"""Export service: export translation results to various formats."""

import json
import logging
from pathlib import Path
from typing import Optional

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from sqlalchemy.orm import Session

from ruzh_translator.models.segment import AlignmentPair, Segment
from ruzh_translator.utils.tmx_parser import export_tmx

logger = logging.getLogger(__name__)


def export_to_xlsx(
    pairs: list,
    output_path: str,
    source_label: str = "俄文",
    target_label: str = "中文",
    include_metadata: bool = True,
) -> str:
    """Export alignment pairs to Excel format.

    Args:
        pairs: List of AlignmentPair instances or dicts with source_text/target_text.
        output_path: Output file path.
        source_label: Column label for source language.
        target_label: Column label for target language.
        include_metadata: Whether to include status and confidence columns.

    Returns:
        Output file path.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "翻译结果"

    # Header styling
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # Columns
    if include_metadata:
        headers = ["序号", source_label, target_label, "置信度", "状态", "段落"]
    else:
        headers = ["序号", source_label, target_label]

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_align

    # Data rows
    wrap_align = Alignment(vertical="top", wrap_text=True)
    for i, pair in enumerate(pairs):
        row = i + 2
        if isinstance(pair, dict):
            src = pair.get("source_text", "")
            tgt = pair.get("target_text", "")
            conf = pair.get("confidence_score", pair.get("confidence", 0))
            status = pair.get("status", "")
            para = pair.get("paragraph_index", pair.get("para_index", 0))
        else:
            src = pair.source_text or ""
            tgt = pair.target_text or ""
            conf = pair.confidence_score or 0
            status = getattr(pair, "status", "")
            para = pair.paragraph_index or 0

        ws.cell(row=row, column=1, value=i + 1).alignment = header_align
        ws.cell(row=row, column=2, value=src).alignment = wrap_align
        ws.cell(row=row, column=3, value=tgt).alignment = wrap_align
        if include_metadata:
            ws.cell(row=row, column=4, value=round(conf, 4) if conf else 0).alignment = header_align
            ws.cell(row=row, column=5, value=status).alignment = header_align
            ws.cell(row=row, column=6, value=para + 1 if isinstance(para, int) else 0).alignment = header_align

    # Column widths
    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 60
    ws.column_dimensions["C"].width = 60
    if include_metadata:
        ws.column_dimensions["D"].width = 10
        ws.column_dimensions["E"].width = 12
        ws.column_dimensions["F"].width = 8

    # Freeze header
    ws.freeze_panes = "A2"

    wb.save(output_path)
    return output_path


def export_to_docx(
    pairs: list,
    output_path: str,
    source_label: str = "俄语原文",
    target_label: str = "中文译文",
) -> str:
    """Export alignment pairs to a bilingual Word document.

    Creates a two-column table: source on left, target on right.

    Args:
        pairs: List of pairs.
        output_path: Output path.
        source_label: Source column header.
        target_label: Target column header.

    Returns:
        Output file path.
    """
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)

    # Title
    title = doc.add_heading("双语对齐翻译结果", level=1)

    # Table
    table = doc.add_table(rows=1, cols=2, style="Table Grid")
    table.autofit = True

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = source_label
    hdr_cells[1].text = target_label
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for pair in pairs:
        if isinstance(pair, dict):
            src = pair.get("source_text", "")
            tgt = pair.get("target_text", "")
        else:
            src = pair.source_text or ""
            tgt = pair.target_text or ""

        row = table.add_row()
        row.cells[0].text = src
        row.cells[1].text = tgt

    doc.save(output_path)
    return output_path


def export_to_html(
    pairs: list,
    output_path: str,
    source_label: str = "俄语",
    target_label: str = "中文",
) -> str:
    """Export alignment pairs to an HTML preview page.

    Args:
        pairs: List of pairs.
        output_path: Output path.
        source_label: Source column label.
        target_label: Target column label.

    Returns:
        Output file path.
    """
    rows_html = ""
    for i, pair in enumerate(pairs):
        if isinstance(pair, dict):
            src = pair.get("source_text", "").replace("<", "&lt;").replace(">", "&gt;")
            tgt = pair.get("target_text", "").replace("<", "&lt;").replace(">", "&gt;")
            conf = pair.get("confidence_score", pair.get("confidence", 0))
        else:
            src = (pair.source_text or "").replace("<", "&lt;").replace(">", "&gt;")
            tgt = (pair.target_text or "").replace("<", "&lt;").replace(">", "&gt;")
            conf = pair.confidence_score or 0

        conf_color = (
            "green" if conf > 0.8 else ("orange" if conf > 0.5 else "red")
        )

        rows_html += f"""
        <tr>
            <td class="num">{i + 1}</td>
            <td class="source">{src}</td>
            <td class="target">{tgt}</td>
            <td class="conf" style="color:{conf_color}">{conf:.2%}</td>
        </tr>"""

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>双语对齐翻译</title>
    <style>
        body {{ font-family: -apple-system, 'Segoe UI', sans-serif; margin: 20px; background: #f5f5f5; }}
        h1 {{ color: #333; text-align: center; }}
        table {{ width: 100%; border-collapse: collapse; background: white; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }}
        th {{ background: #4472C4; color: white; padding: 12px; text-align: left; position: sticky; top: 0; }}
        td {{ padding: 10px; border-bottom: 1px solid #eee; vertical-align: top; }}
        .num {{ width: 50px; text-align: center; color: #999; }}
        .source {{ width: 45%; }}
        .target {{ width: 45%; }}
        .conf {{ width: 60px; text-align: center; font-weight: bold; }}
        tr:hover {{ background: #f0f4ff; }}
    </style>
</head>
<body>
    <h1>{source_label} → {target_label} 双语对齐</h1>
    <table>
        <thead>
            <tr>
                <th>#</th>
                <th>{source_label}</th>
                <th>{target_label}</th>
                <th>置信度</th>
            </tr>
        </thead>
        <tbody>
            {rows_html}
        </tbody>
    </table>
</body>
</html>"""

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html.strip())

    return output_path


def export_to_tmx(
    session: Session,
    project_id: str,
    output_path: str,
    source_lang: str = "ru",
    target_lang: str = "zh-CN",
    only_approved: bool = False,
) -> str:
    """Export project translation pairs to TMX format.

    Args:
        session: Database session.
        project_id: Project ID.
        output_path: Output path.
        source_lang: Source language code.
        target_lang: Target language code.
        only_approved: If True, only export approved segments.

    Returns:
        Output file path.
    """
    query = session.query(AlignmentPair).filter(
        AlignmentPair.project_id == project_id
    )

    if only_approved:
        # Join with segments to filter by status
        query = query.join(
            Segment,
            AlignmentPair.source_segment_id == Segment.id,
        ).filter(Segment.status == "approved")

    pairs = query.all()

    pair_dicts = [
        {"source_text": p.source_text or "", "target_text": p.target_text or ""}
        for p in pairs
        if p.target_text  # Only export pairs that have translations
    ]

    export_tmx(pair_dicts, output_path, source_lang, target_lang)
    return output_path


def export_project(
    session: Session,
    project_id: str,
    output_dir: str,
    formats: list[str] = None,
) -> dict:
    """Export a project to multiple formats at once.

    Queries BOTH AlignmentPair and Segment tables to handle:
    - Aligned translations (AlignmentPair)
    - Direct translations (Segment with target_text, no AlignmentPair)

    Args:
        session: Database session.
        project_id: Project ID.
        output_dir: Output directory.
        formats: List of format names ('tmx', 'xlsx', 'docx', 'html').

    Returns:
        Dict mapping format to output file path.
    """
    if formats is None:
        formats = ["xlsx", "tmx"]

    from ruzh_translator.models.project import Project
    from ruzh_translator.models.segment import Segment

    project = session.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise ValueError(f"Project not found: {project_id}")

    # Query both data sources
    alignment_pairs = (
        session.query(AlignmentPair)
        .filter(AlignmentPair.project_id == project_id)
        .order_by(AlignmentPair.paragraph_index, AlignmentPair.pair_index)
        .all()
    )

    # Also get direct-translated Segments (those with target_text but no alignment pair)
    # We identify them by: Segments that have translations
    direct_segments = (
        session.query(Segment)
        .filter(
            Segment.project_id == project_id,
            Segment.target_text != "",
            Segment.target_text.isnot(None),
        )
        .order_by(Segment.paragraph_index, Segment.segment_index)
        .all()
    )

    # Merge: use alignment pairs if available, otherwise fall back to segments
    if alignment_pairs:
        pairs = [
            {
                "source_text": p.source_text or "",
                "target_text": p.target_text or "",
                "confidence_score": p.confidence_score,
                "paragraph_index": p.paragraph_index,
                "status": getattr(
                    session.query(Segment).filter(Segment.id == p.source_segment_id).first(),
                    "status", ""
                ) if p.source_segment_id else "",
            }
            for p in alignment_pairs
        ]
    elif direct_segments:
        # Use direct segments, converting to the dict format expected by exporters
        pairs = [
            {
                "source_text": seg.source_text or "",
                "target_text": seg.target_text or "",
                "confidence_score": 1.0,
                "paragraph_index": seg.paragraph_index,
                "status": seg.status or "",
            }
            for seg in direct_segments
        ]
    else:
        pairs = []

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    base_name = project.name.replace(" ", "_")
    results = {}

    if not pairs:
        # Still create files but warn
        import logging
        logging.warning(f"No translation data to export for project {project_id}")

    if "xlsx" in formats:
        path = output_dir / f"{base_name}.xlsx"
        export_to_xlsx(pairs, str(path))
        results["xlsx"] = str(path)

    if "docx" in formats:
        path = output_dir / f"{base_name}_bilingual.docx"
        export_to_docx(pairs, str(path))
        results["docx"] = str(path)

    if "html" in formats:
        path = output_dir / f"{base_name}.html"
        export_to_html(pairs, str(path))
        results["html"] = str(path)

    if "tmx" in formats:
        path = output_dir / f"{base_name}.tmx"
        if pairs:
            # Convert merged dicts to TMX-compatible format
            tmx_pairs = [
                {"source_text": p["source_text"], "target_text": p["target_text"]}
                for p in pairs if p["target_text"]
            ]
            export_tmx(
                tmx_pairs, str(path),
                project.source_lang, project.target_lang,
            )
        else:
            export_tmx(
                [], str(path),
                project.source_lang, project.target_lang,
            )
        results["tmx"] = str(path)

    return results
